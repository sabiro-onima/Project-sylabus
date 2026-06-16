"""
Export service – generuje PDF i DOCX dla sylabusów.
  PDF  → WeasyPrint (HTML → PDF)
  DOCX → python-docx
"""
import io
import uuid
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Environment, BaseLoader
from weasyprint import HTML

from app.models.syllabus import SyllabusVersion


# ─── HTML TEMPLATE FOR PDF ───────────────────────────────────────────────────

PDF_TEMPLATE = """
<!DOCTYPE html>
<html lang="pl">
<head>
<meta charset="UTF-8">
<style>
  body { font-family: DejaVu Sans, Arial, sans-serif; font-size: 11pt; color: #1a1a1a; margin: 2cm; }
  h1 { font-size: 18pt; color: #1a3a6b; border-bottom: 2px solid #1a3a6b; padding-bottom: 8px; }
  h2 { font-size: 13pt; color: #1a3a6b; margin-top: 20px; border-left: 4px solid #1a3a6b; padding-left: 8px; }
  .meta-table { width: 100%; border-collapse: collapse; margin: 16px 0; }
  .meta-table td { padding: 6px 10px; border: 1px solid #ddd; }
  .meta-table td:first-child { font-weight: bold; background: #f4f6fb; width: 35%; }
  .hours-table { width: 100%; border-collapse: collapse; margin: 10px 0; }
  .hours-table th { background: #1a3a6b; color: white; padding: 6px 10px; text-align: left; }
  .hours-table td { padding: 6px 10px; border: 1px solid #ddd; }
  .hours-table tr:nth-child(even) td { background: #f4f6fb; }
  .outcome { margin: 6px 0; padding: 6px 10px; background: #f4f6fb; border-radius: 4px; }
  .outcome-code { font-weight: bold; color: #1a3a6b; }
  .biblio-item { margin: 4px 0; padding-left: 16px; }
  .badge { display: inline-block; padding: 2px 8px; border-radius: 10px; font-size: 9pt; }
  .badge-primary { background: #1a3a6b; color: white; }
  .badge-supplementary { background: #6b8cba; color: white; }
  .footer { margin-top: 30px; font-size: 9pt; color: #888; border-top: 1px solid #ddd; padding-top: 8px; }
  p { line-height: 1.6; }
</style>
</head>
<body>

<h1>{{ version.title_pl }}</h1>
{% if version.title_en %}<p style="color:#666;margin-top:-10px;font-style:italic">{{ version.title_en }}</p>{% endif %}

<h2>Informacje ogólne</h2>
<table class="meta-table">
  <tr><td>Kod przedmiotu</td><td>{{ course_code }}</td></tr>
  <tr><td>Rok akademicki</td><td>{{ version.academic_year }}</td></tr>
  <tr><td>Forma zajęć</td><td>{{ version.course_type }}</td></tr>
  <tr><td>Semestr</td><td>{{ version.semester }} (sem. {{ version.semester_number }})</td></tr>
  <tr><td>Punkty ECTS</td><td>{{ version.ects_credits }}</td></tr>
  <tr><td>Język</td><td>{{ version.language }}</td></tr>
</table>

{% if version.description %}
<h2>Opis przedmiotu</h2>
<p>{{ version.description }}</p>
{% endif %}

{% if version.learning_objectives %}
<h2>Cele kształcenia</h2>
<p>{{ version.learning_objectives }}</p>
{% endif %}

{% if version.prerequisites %}
<h2>Wymagania wstępne</h2>
<p>{{ version.prerequisites }}</p>
{% endif %}

{% if version.learning_outcomes %}
<h2>Efekty kształcenia</h2>
{% for eo in version.learning_outcomes %}
<div class="outcome">
  <span class="outcome-code">{{ eo.code }}</span> [{{ eo.category }}] — {{ eo.description }}
</div>
{% endfor %}
{% endif %}

<h2>Liczba godzin</h2>
<table class="hours-table">
  <tr><th>Forma</th><th>Godziny</th></tr>
  <tr><td>Wykład</td><td>{{ version.hours_lecture }}</td></tr>
  <tr><td>Laboratorium</td><td>{{ version.hours_laboratory }}</td></tr>
  <tr><td>Ćwiczenia</td><td>{{ version.hours_exercise }}</td></tr>
  <tr><td>Seminarium</td><td>{{ version.hours_seminar }}</td></tr>
  <tr><td>Projekt</td><td>{{ version.hours_project }}</td></tr>
  <tr><td>Samodzielna nauka</td><td>{{ version.hours_self_study }}</td></tr>
  <tr><td><strong>ŁĄCZNIE</strong></td><td><strong>{{ total_hours }}</strong></td></tr>
</table>

{% if version.assessment_methods %}
<h2>Metody oceniania</h2>
<table class="hours-table">
  <tr><th>Metoda</th><th>Waga</th><th>Opis</th></tr>
  {% for am in version.assessment_methods %}
  <tr>
    <td>{{ am.method }}</td>
    <td>{{ am.weight }}%</td>
    <td>{{ am.description or '—' }}</td>
  </tr>
  {% endfor %}
</table>
{% endif %}

{% if version.bibliography %}
<h2>Literatura</h2>
{% for b in version.bibliography %}
<div class="biblio-item">
  <span class="badge badge-{{ b.type }}">{{ b.type }}</span>
  {{ b.citation }}
  {% if b.url %} — <a href="{{ b.url }}">{{ b.url }}</a>{% endif %}
</div>
{% endfor %}
{% endif %}

<div class="footer">
  Wygenerowano: {{ generated_at }} | Wersja {{ version.version_number }} | Status: {{ version.status }}
</div>

</body>
</html>
"""


class ExportService:

    # ─── PDF ─────────────────────────────────────────────────────────────────

    def generate_pdf(self, version: SyllabusVersion, course_code: str) -> bytes:
        """Generuje PDF syllabusa jako bytes."""
        env = Environment(loader=BaseLoader())
        template = env.from_string(PDF_TEMPLATE)

        html_content = template.render(
            version=version,
            course_code=course_code,
            total_hours=self._total_hours(version),
            generated_at=datetime.now().strftime("%d.%m.%Y %H:%M"),
        )

        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes

    # ─── DOCX ────────────────────────────────────────────────────────────────

    def generate_docx(self, version: SyllabusVersion, course_code: str) -> bytes:
        """Generuje DOCX syllabusa jako bytes."""
        doc = Document()

        # ── Styl dokumentu ──────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Tytuł ────────────────────────────────────────────────────────────
        title = doc.add_heading(version.title_pl, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if version.title_en:
            sub = doc.add_paragraph(version.title_en)
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].italic = True

        doc.add_paragraph()

        # ── Informacje ogólne ─────────────────────────────────────────────────
        doc.add_heading("Informacje ogólne", level=2)
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"

        info_rows = [
            ("Kod przedmiotu", course_code),
            ("Rok akademicki", version.academic_year),
            ("Forma zajęć", version.course_type),
            ("Semestr", f"{version.semester} (sem. {version.semester_number})"),
            ("Punkty ECTS", str(version.ects_credits)),
            ("Język", version.language),
        ]
        for i, (label, value) in enumerate(info_rows):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value

        # ── Sekcje tekstowe ────────────────────────────────────────────────
        sections = [
            ("Opis przedmiotu", version.description),
            ("Cele kształcenia", version.learning_objectives),
            ("Wymagania wstępne", version.prerequisites),
        ]
        for heading, content in sections:
            if content:
                doc.add_heading(heading, level=2)
                doc.add_paragraph(content)

        # ── Efekty kształcenia ────────────────────────────────────────────
        if version.learning_outcomes:
            doc.add_heading("Efekty kształcenia", level=2)
            for eo in version.learning_outcomes:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{eo['code']} [{eo['category']}]: ").bold = True
                p.add_run(eo["description"])

        # ── Godziny ──────────────────────────────────────────────────────────
        doc.add_heading("Liczba godzin", level=2)
        hours_table = doc.add_table(rows=7, cols=2)
        hours_table.style = "Table Grid"
        hours_data = [
            ("Wykład", version.hours_lecture),
            ("Laboratorium", version.hours_laboratory),
            ("Ćwiczenia", version.hours_exercise),
            ("Seminarium", version.hours_seminar),
            ("Projekt", version.hours_project),
            ("Samodzielna nauka", version.hours_self_study),
            ("ŁĄCZNIE", self._total_hours(version)),
        ]
        for i, (label, hours) in enumerate(hours_data):
            row = hours_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(hours)
            if label == "ŁĄCZNIE":
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].bold = True

        # ── Metody oceniania ──────────────────────────────────────────────
        if version.assessment_methods:
            doc.add_heading("Metody oceniania", level=2)
            am_table = doc.add_table(rows=len(version.assessment_methods) + 1, cols=3)
            am_table.style = "Table Grid"
            headers = am_table.rows[0]
            for i, h in enumerate(["Metoda", "Waga", "Opis"]):
                headers.cells[i].text = h
                headers.cells[i].paragraphs[0].runs[0].bold = True
            for i, am in enumerate(version.assessment_methods, 1):
                row = am_table.rows[i]
                row.cells[0].text = am["method"]
                row.cells[1].text = f"{am['weight']}%"
                row.cells[2].text = am.get("description") or "—"

        # ── Literatura ────────────────────────────────────────────────────
        if version.bibliography:
            doc.add_heading("Literatura", level=2)
            for b in version.bibliography:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{b['type']}] ").bold = True
                p.add_run(b["citation"])
                if b.get("url"):
                    p.add_run(f" — {b['url']}")

        # ── Stopka ────────────────────────────────────────────────────────
        doc.add_paragraph()
        footer = doc.add_paragraph(
            f"Wygenerowano: {datetime.now().strftime('%d.%m.%Y %H:%M')} | "
            f"Wersja {version.version_number} | Status: {version.status}"
        )
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # ── Zapis do bytes ────────────────────────────────────────────────
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ─── HELPERS ─────────────────────────────────────────────────────────────

    def _total_hours(self, v: SyllabusVersion) -> int:
        return (
            v.hours_lecture + v.hours_laboratory + v.hours_exercise
            + v.hours_seminar + v.hours_project + v.hours_self_study
        )
